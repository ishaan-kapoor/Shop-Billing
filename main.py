import datetime
import os
from tkinter import (
    Button,
    Entry,
    Frame,
    IntVar,
    Label,
    StringVar,
    TclError,
    Tk,
    messagebox,
)

import docx
from tabulate import tabulate


class CustomEntry(Entry):
    def __init__(self, master=None, textvariable=None, background="#121212", width=15):
        super().__init__(
            master,
            bd=0,
            width=width,
            fg="white",
            justify="center",
            font=("Times", 18),
            highlightthickness=0,
            background=background,
            selectbackground="white",
            textvariable=textvariable,
            selectforeground="#464646",
        )

    def grid(self, row=0, column=0, pad=2):
        super().grid(row=row, column=column, padx=pad, pady=pad)


class CustomButton(Button):
    def __init__(self, master=None, text="", command=None):
        super().__init__(
            master,
            text=text,
            fg="white",
            command=command,
            background="#a6a6a6".replace("a", "4"),
        )

    def grid(self, row=0, column=0, pad=2):
        super().grid(row=row, column=column, padx=pad, pady=pad)

    def pack(self, side="Top", pad=2):
        super().pack(side=side, padx=pad, pady=pad)


class CustomLabel(Label):
    def __init__(self, master=None, text=None, textvariable=None):
        super().__init__(
            master,
            bd=0,
            text=text,
            fg="white",
            justify="center",
            font=("Times", 18),
            background="#121212",
            highlightthickness=0,
            highlightcolor="#464646",
            textvariable=textvariable,
            highlightbackground="white",
        )

    def grid(self, row=0, column=0, pad=2):
        super().grid(row=row, column=column, padx=pad, pady=pad)


class CustomFrame(Frame):
    def __init__(self, master=None):
        super().__init__(master, background="#121212", borderwidth=0)

    def pack(self):
        super().pack(pady=10, padx=10, fill="both")


validation_int = lambda entry: entry.isdigit() or (entry == "")


def validation_decimal(entry):
    if validation_int(entry):
        return True
    for char in entry:
        if char not in ".0123456789":
            return False
    else:
        return True


def validation_discount(entry):
    if validation_decimal(entry):
        return True
    for char in entry:
        if char not in ".0123456789%":
            return False
    else:
        return True


def getDate():
    return datetime.datetime.now().strftime(r"%I:%M %p    %d %B, %Y")


def newItem():
    if not running:
        return

    global count
    count += 1

    def delete():
        if not running:
            return

        global count
        if count == 1:
            return clear()
        count -= 1
        textVariables.remove(variables)
        widgets.remove(items)
        sl_no_widget.destroy()
        desc_widget.destroy()
        qty_widget.destroy()
        rate_widget.destroy()
        amt_widget.destroy()
        clr_button.destroy()
        del_button.destroy()
        for i, (sl_no, *_) in enumerate(textVariables, start=1):
            sl_no.set(i)
        for i, (sl_no, desc, qty, rate, amt, clr_bt, del_bt) in enumerate(
            widgets, start=3
        ):
            sl_no.grid(row=i, column=0)
            desc.grid(row=i, column=1)
            qty.grid(row=i, column=2)
            rate.grid(row=i, column=3)
            amt.grid(row=i, column=4)
            clr_bt.grid(row=i, column=5)
            del_bt.grid(row=i, column=6)
        total()

    def clear():
        if not running:
            return
        desc_var.set("-")
        qty_var.set("1.0")
        rate_var.set("1")

    sl_no_var = IntVar(value=count)
    desc_var = StringVar(value="-")
    qty_var = StringVar(value="1.0")
    rate_var = StringVar(value="1")
    amt_var = IntVar(value=0)
    variables = (sl_no_var, desc_var, qty_var, rate_var, amt_var)
    textVariables.insert(count - 1, variables)
    for var in variables[:-1]:
        var.trace_add("write", total)

    sl_no_widget = CustomLabel(products, textvariable=sl_no_var)
    desc_widget = CustomEntry(products, textvariable=desc_var)
    qty_widget = CustomEntry(products, textvariable=qty_var)
    rate_widget = CustomEntry(products, textvariable=rate_var)
    amt_widget = CustomLabel(products, textvariable=amt_var)
    clr_button = CustomButton(products, text="Clear", command=clear)
    del_button = CustomButton(products, text="Delete", command=delete)
    items = (
        sl_no_widget,
        desc_widget,
        qty_widget,
        rate_widget,
        amt_widget,
        clr_button,
        del_button,
    )
    widgets.insert(count - 1, items)
    qty_widget.config(validate="key", validatecommand=(decimal_validation, r"%P"))
    rate_widget.config(validate="key", validatecommand=(int_validation, r"%P"))

    sl_no_widget.grid(row=count + 2, column=0)
    desc_widget.grid(row=count + 2, column=1)
    qty_widget.grid(row=count + 2, column=2)
    rate_widget.grid(row=count + 2, column=3)
    amt_widget.grid(row=count + 2, column=4)
    clr_button.grid(row=count + 2, column=5)
    del_button.grid(row=count + 2, column=6)

    total()


def delLastItem():
    if not running:
        return

    global count
    if count <= 0:
        return
    elif count == 1:
        temp = textVariables[0]
        temp[1].set("-")
        temp[2].set("1.0")
        temp[3].set("1")
    elif count > 1:
        textVariables.pop()
        for widget in widgets.pop():
            widget.destroy()
        count -= 1
    total()


def reset():
    if not running:
        return
    for _ in range(count):
        delLastItem()


def toggleDiscount():
    global showDiscount
    showDiscount = not showDiscount
    discountVar.set("")
    if showDiscount:
        totalLabel.grid(row=1, column=1)
        totalValLabel.grid(row=1, column=2)
        discountLabel.grid(row=2, column=1)
        discountEntry.grid(row=2, column=2)
    else:
        discountLabel.grid_forget()
        discountEntry.grid_forget()
        totalLabel.grid_forget()
        totalValLabel.grid_forget()


def total(*args):
    if not running:
        return
    root.update()
    date.set(getDate())
    val = 0
    if count <= 0:
        newItem()
    for *_, qty, rate, amt in textVariables:
        for var in (qty, rate):
            try:
                var.set(var.get().strip().lstrip("0"))
            except TclError:
                pass
        try:
            amt.set(
                temp := 0
                if (qty.get() == "") or (rate.get() == "")
                else round(float(qty.get()) * int(rate.get()))
            )
        except TclError:
            temp = 0
        val += temp
    totalVar.set(val)
    try:
        discountVar.set(discountVar.get().lstrip("0"))
    except TclError:
        pass
    if (temp := discountVar.get()) in ("", "."):
        disc = 0
    elif "%" in temp:
        disc = val * eval(temp.replace("%", "")) / 100
    else:
        disc = eval(temp)
    grandTotalVar.set(round(val - disc))


def close():
    global running
    if not running:
        return
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        running = False
        root.destroy()


def getBillNumber():
    path = ""
    temp = str(date.get()).split("    ")[-1].split()
    for folder in ("Receipts", temp[2], temp[1][:-1], temp[0]):
        path = os.path.join(path, folder)
        if not os.path.exists(path):
            return "1"
    if not (temp := os.listdir(path)):
        return "1"
    return str(int(max(temp, key=lambda x: int(x[:-4]))[:-4]) + 1)


def generateReceipt():
    if not running:
        return
    generateReceipt_Word()
    generateReceipt_Text()
    billNumber.set(str(int(billNumber.get()) + 1))


def generateReceipt_Word():
    if not running:
        return
    doc = docx.Document("Format.docx")
    table = doc.tables[0]
    for variables in textVariables:
        row_cells = table.add_row().cells
        row_cells[0].text = str(variables[0].get())
        row_cells[1].text = str(variables[1].get())
        row_cells[2].text = str(variables[2].get())
        row_cells[3].text = str(variables[3].get())
        row_cells[4].text = str(variables[4].get())
    table.add_row()
    if (temp_discount := discountVar.get()) not in ("", "0"):
        row_cells = table.add_row().cells
        row_cells[3].text = "Total"
        row_cells[4].text = str(totalVar.get())
        row_cells = table.add_row().cells
        row_cells[3].text = "Discount"
        row_cells[4].text = temp_discount
    row_cells = table.add_row().cells
    row_cells[3].text = "Grand Total"
    row_cells[4].text = str(grandTotalVar.get())

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = (
                    docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.size = docx.shared.Pt(8)

    para = doc.paragraphs[-1]
    if (temp := customer.get()).strip() not in ("", "-"):
        para.add_run(f"Customer Details: \t{temp}\n").font.size = docx.shared.Pt(8)
    if (temp := dealer.get()).strip() not in ("", "-"):
        para.add_run(f"Dealer: \t{temp}\n").font.size = docx.shared.Pt(8)
    if (temp := paymentMode.get()).strip() not in ("", "-"):
        para.add_run(f"Payment Mode: \t{temp}\n").font.size = docx.shared.Pt(8)

    doc.save("Receipt.docx")
    os.startfile("Receipt.docx")


def generateReceipt_Text():
    if not running:
        return

    data = [
        (),
        ("Sl No", "Description", "Quantity", "Rate", "Amount"),
        *(tuple(j.get() for j in i) for i in textVariables),
        (),
    ]

    temp_customer = (
        ()
        if (temp := customer.get()).strip() in ("", "-")
        else ("Customer Details", temp)
    )
    temp_dealer = (
        () if (temp := dealer.get()).strip() in ("", "-") else ("Dealer", temp)
    )
    temp_payment = (
        ()
        if (temp := paymentMode.get()).strip() in ("", "-")
        else ("Payment Mode", temp)
    )

    if (temp_discount := discountVar.get()) not in ("", "0"):
        data += [
            ("", "", "", "Total", totalVar.get()),
            ("", "", "", "Discount", temp_discount),
        ]
    data.append(("", "", "", "Grand Total", grandTotalVar.get()))

    number = billNumber.get()
    receipt = tabulate(data, tablefmt="plain", stralign="center", numalign="left")
    temp = (f"Bill Number: {number}", date.get())
    spaces = len(receipt.split("\n", 2)[1]) - sum(map(len, temp))
    temp = temp[0] + " " * spaces + temp[1] + "\n"
    receipt = (
        temp
        + receipt
        + "\n"
        + tabulate((temp_customer, temp_dealer, temp_payment), tablefmt="plain").strip()
    )

    path = ""
    temp = str(date.get()).split("    ")[-1].split()
    for folder in ("Receipts", temp[2], temp[1][:-1], temp[0]):
        path = os.path.join(path, folder)
        if not os.path.exists(path):
            os.mkdir(path)
    if not number.isdigit():
        number = getBillNumber()
    path = os.path.join(path, number + ".txt")
    with open(path, "wt") as file:
        file.write(receipt)
    os.startfile(path)
    return receipt


count = 0
widgets = []
running = True
textVariables = []
showDiscount = False

root = Tk()
root.title("UTSAV Billing - By Ishaan")
root.config(bg="#223441")
root.resizable(width=True, height=True)
root.protocol("WM_DELETE_WINDOW", close)

int_validation = root.register(validation_int)
decimal_validation = root.register(validation_decimal)
discount_validation = root.register(validation_discount)


header = CustomFrame(root)
CustomButton(header, text="Delete Last Item", command=delLastItem).pack(side="left")
CustomButton(header, text="Reset", command=reset).pack(side="left")
CustomButton(header, text="Toggle Discount", command=toggleDiscount).pack(side="left")
CustomButton(header, text="Receipt", command=generateReceipt).pack(side="left")
CustomButton(header, text="Add Item", command=newItem).pack(side="left")
date = StringVar(value=getDate())
CustomLabel(header, textvariable=date).pack(side="right")
header.pack()


products = CustomFrame(root)
billNumber = StringVar(value=getBillNumber())
CustomLabel(products, text="Bill Number: ").grid(row=1, column=0)
CustomEntry(products, textvariable=billNumber, background="#464646").grid(
    row=1, column=1
)
CustomLabel(products, text="Sl No").grid(row=2, column=0, pad=10)
CustomLabel(products, text="Description").grid(row=2, column=1, pad=10)
CustomLabel(products, text="Quantity").grid(row=2, column=2, pad=10)
CustomLabel(products, text="Rate").grid(row=2, column=3, pad=10)
CustomLabel(products, text="Amount").grid(row=2, column=4, pad=10)
products.pack()


totaling = CustomFrame(root)
totaling.columnconfigure(0, weight=86)
totaling.columnconfigure(1, weight=1)
totaling.columnconfigure(2, weight=1)
totaling.columnconfigure(3, weight=12)
totalVar = IntVar()
totalLabel = CustomLabel(totaling, text="Total: ")
if showDiscount:
    totalLabel.grid(row=1, column=1)
totalValLabel = CustomLabel(totaling, textvariable=totalVar)
if showDiscount:
    totalValLabel.grid(row=1, column=2)
discountVar = StringVar(value="")
discountVar.trace_add("write", total)
discountLabel = CustomLabel(totaling, text="Discount: ")
if showDiscount:
    discountLabel.grid(row=2, column=1)
discountEntry = CustomEntry(totaling, textvariable=discountVar, background="#464646")
discountEntry.config(validate="key", validatecommand=(discount_validation, r"%P"))
if showDiscount:
    discountEntry.grid(row=2, column=2)
grandTotalVar = IntVar()
CustomLabel(totaling, text="Grand Total: ").grid(row=3, column=1)
CustomLabel(totaling, textvariable=grandTotalVar).grid(row=3, column=2)
totaling.pack()

details = CustomFrame(root)
customer = StringVar(value="-")
CustomLabel(details, text="Customer Details").grid(row=1, column=1)
CustomEntry(details, textvariable=customer, background="#464646").grid(row=1, column=2)
dealer = StringVar(value="-")
CustomLabel(details, text="Dealer").grid(row=2, column=1)
CustomEntry(details, textvariable=dealer, background="#464646").grid(row=2, column=2)
paymentMode = StringVar(value="CASH")
CustomLabel(details, text="Payment Mode").grid(row=3, column=1)
CustomEntry(details, textvariable=paymentMode, background="#464646").grid(
    row=3, column=2
)
details.pack()

footer = CustomFrame(root)
footer.pack()

newItem()

root.mainloop()
exit()
